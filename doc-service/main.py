import os
import re
import hashlib
import subprocess
from datetime import datetime

import psycopg2
from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = FastAPI()

DB_CONFIG = {
    "host":     os.environ.get("DB_HOST", "postgres"),
    "port":     5432,
    "dbname":   os.environ.get("DB_NAME", "n8n"),
    "user":     os.environ.get("DB_USER", "n8n"),
    "password": os.environ.get("DB_PASSWORD", ""),
}

DOCS_PATH = "/documents"


class DocRequest(BaseModel):
    reference: str
    document_type: str
    content: str
    parties: dict
    compliance_flags: list = []
    attorney_review_count: int = 0
    approved: bool = False


def sha256_file(path: str) -> str:
    with open(path, "rb") as f:
        return hashlib.sha256(f.read()).hexdigest()


def log_to_db(req: DocRequest, doc_hash: str):
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur  = conn.cursor()
        cur.execute("""
            INSERT INTO document_log
                (reference, document_type, party1, party2,
                 doc_hash, llm_model, compliance_flags, approved, created_at)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, NOW())
            ON CONFLICT (reference) DO UPDATE
                SET doc_hash = EXCLUDED.doc_hash,
                    approved = EXCLUDED.approved
        """, (
            req.reference,
            req.document_type,
            list(req.parties.values())[0],
            list(req.parties.values())[1] if len(req.parties) > 1 else None,
            doc_hash,
            os.environ.get("OLLAMA_MODEL", "mistral:latest"),
            req.compliance_flags,
            req.approved,
        ))
        conn.commit()
        cur.close()
        conn.close()
    except Exception as e:
        print(f"[WARN] DB logging failed: {e}")


@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/generate")
def generate_document(req: DocRequest):
    os.makedirs(DOCS_PATH, exist_ok=True)

    doc = Document()

    # Title
    title = doc.add_heading(req.document_type.upper(), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    doc.add_paragraph(f"Reference:  {req.reference}")
    doc.add_paragraph(f"Date:       {datetime.today().strftime('%d %B %Y')}")
    doc.add_paragraph(f"Parties:    {' / '.join(req.parties.values())}")
    doc.add_paragraph("")

    # Draft watermark
    if not req.approved:
        p   = doc.add_paragraph("⚠  DRAFT — NOT LEGALLY BINDING — PENDING ATTORNEY REVIEW")
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        doc.add_paragraph("")

    # Compliance warnings
    if req.compliance_flags:
        doc.add_paragraph("COMPLIANCE WARNINGS:")
        for flag in req.compliance_flags:
            doc.add_paragraph(f"  • {flag}")
        doc.add_paragraph("")

    # Document body
    for para in req.content.split("\n\n"):
        stripped = para.strip()
        if stripped:
            doc.add_paragraph(stripped)

    # Signature block
    doc.add_paragraph("\n\n")
    party_names = list(req.parties.values())
    doc.add_paragraph("_" * 30 + "          " + "_" * 30)
    doc.add_paragraph(
        f"{party_names[0]}"
        + "          "
        + (party_names[1] if len(party_names) > 1 else "")
    )
    doc.add_paragraph("Date: _______________          Date: _______________")

    # POPIA notice
    doc.add_paragraph("")
    popia = doc.add_paragraph(
        "POPIA NOTICE: Personal information contained herein is processed "
        "in accordance with the Protection of Personal Information Act 4 of 2013. "
        "This information will not be disclosed to third parties without prior consent."
    )
    popia.runs[0].font.size = Pt(8)

    # Save .docx
    docx_path = f"{DOCS_PATH}/{req.reference}.docx"
    pdf_path  = f"{DOCS_PATH}/{req.reference}.pdf"
    doc.save(docx_path)

    # Convert to PDF
    result = subprocess.run([
        "libreoffice", "--headless", "--convert-to", "pdf",
        "--outdir", DOCS_PATH, docx_path
    ], capture_output=True, text=True)

    if result.returncode != 0:
        raise HTTPException(
            status_code=500,
            detail=f"PDF conversion failed: {result.stderr}"
        )

    doc_hash = sha256_file(pdf_path)
    log_to_db(req, doc_hash)

    return {
        "reference": req.reference,
        "docx_path": docx_path,
        "pdf_path":  pdf_path,
        "hash":      doc_hash
    }
